import fitz
import logging
from docx import Document
from io import BytesIO
from typing import Tuple

logger = logging.getLogger("document-parser")


def parse_pdf_bytes(file_bytes: bytes) -> str:
    """从PDF字节流中提取文本"""
    try:
        logger.info(f"[PDF解析] 开始解析，文件大小：{len(file_bytes)}字节")
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()

        if not full_text.strip():
            logger.warning("[PDF解析] 解析结果为空，无有效文本内容")
            raise RuntimeError("PDF解析结果为空，可能是图片扫描版文档")

        logger.info(f"[PDF解析] 解析完成，文本长度：{len(full_text.strip())}字符")
        return full_text.strip()

    except RuntimeError:
        # 已经是业务异常，直接抛出
        raise
    except Exception as e:
        logger.error(f"[PDF解析] 解析失败：{str(e)}", exc_info=True)
        raise RuntimeError("PDF文件损坏或无法解析，请检查文件是否正确") from e


def parse_docx_bytes(file_bytes: bytes) -> str:
    """从docx字节流中提取文本"""
    try:
        logger.info(f"[DOCX解析] 开始解析，文件大小：{len(file_bytes)}字节")
        stream = BytesIO(file_bytes)
        doc = Document(stream)

        # 1、提取正文段落
        paragraph_texts = [p.text for p in doc.paragraphs if p.text.strip()]

        # 2、提取表格内容
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_texts.append("|".join(row_cells))

        # 3、合并所有文本
        all_text = paragraph_texts + table_texts
        result = "\n".join(all_text).strip()

        if not result:
            logger.warning("[DOCX解析] 解析结果为空，无有效文本内容")
            raise RuntimeError("文档解析结果为空，无有效文本内容")

        logger.info(f"[DOCX解析] 解析完成，总长度：{len(result)}字符")
        return result

    except RuntimeError:
        # 已经是业务异常，直接抛出
        raise
    except Exception as e:
        logger.error(f"[DOCX解析] 解析失败：{str(e)}", exc_info=True)
        raise RuntimeError("文档解析失败，请检查文件是否损坏或格式是否正确") from e


def parse_document(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    统一入口：根据文件名后缀自动选择解析方式
    返回：(文档类型，解析后的纯文本)
    """
    # 边界校验1：空文件
    if len(file_bytes) == 0:
        logger.warning(f"[文件校验] 上传的文件是空的：{filename}")
        raise RuntimeError("上传的文件是空的，无法解析")

    # 边界校验2：文件大小限制（单文件最大10MB）
    if len(file_bytes) > 10 * 1024 * 1024:
        logger.warning(f"[文件校验] 文件超过10MB限制：{filename}，大小：{len(file_bytes)}字节")
        raise RuntimeError("文件大小超过10MB限制，请上传更小的文档")

    # 边界校验3：文件名合法性校验
    if "." not in filename:
        logger.warning(f"[文件校验] 文件缺少后缀，无法识别类型：{filename}")
        raise RuntimeError("文件缺少后缀，无法识别文件类型")

    suffix = filename.lower().split(".")[-1]

    if suffix == "pdf":
        return "pdf", parse_pdf_bytes(file_bytes)
    elif suffix == "docx":
        return "docx", parse_docx_bytes(file_bytes)
    else:
        logger.warning(f"[文件校验] 不支持的文件格式：{suffix}")
        raise RuntimeError(f"不支持的文件格式：{suffix}，仅支持 pdf 、docx")
